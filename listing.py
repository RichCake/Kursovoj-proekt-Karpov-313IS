import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Путь к проекту
source_folder = "/Users/arsenijkarpov/Documents/Колледж/3 курс/Курсовые/project"

# Каталоги, которые нужно исключить
exclude_dirs = {"venv", "interfaces"}
exclude_files = {"listing.py"}

# Укажите имя выходного Word файла
output_file = "Python_FilesNew.docx"

def add_code_to_document(doc, file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        doc.add_heading(os.path.relpath(file_path, "project"), level=2)  # Название файла как заголовок
        code = file.read()

        # Создаем новый параграф для кода
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(code)

        # Устанавливаем моноширинный шрифт
        run.font.name = 'Courier New'

        # Для совместимости с Word, нужно явно указать стиль шрифта для run
        r = run._element
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Courier New")
        rFonts.set(qn("w:hAnsi"), "Courier New")
        rFonts.set(qn("w:cs"), "Courier New")
        r.insert(0, rFonts)

        # Устанавливаем размер шрифта
        run.font.size = Pt(8)

def main():
    # Создаем документ Word
    doc = Document()
    doc.add_heading("Сборник Python файлов", level=1)

    # Перебираем все файлы .py в указанной папке
    for root, dirs, files in os.walk(source_folder):
        # Исключаем указанные директории
        dirs[:] = [d for d in dirs if d not in exclude_dirs]
        for file in files:
            if file.endswith(".py") and file not in exclude_files:
                file_path = os.path.join(root, file)
                add_code_to_document(doc, file_path)

    # Сохраняем документ
    doc.save(output_file)
    print(f"Документ сохранен как {output_file}")

if __name__ == "__main__":
    main()